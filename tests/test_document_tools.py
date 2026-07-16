"""문서 파일 지원 테스트 — 구형 .xls(xlrd 변환)와 .docx(read_document).

실파일 없이 xlwt/python-docx로 심은 파일(doc_dir)로 검증한다.
"""

import datetime

import pytest
import xlwt
from docx import Document as DocxDocument
from openpyxl import Workbook

from agent.tools.documents import read_document
from agent.tools.excel import (
    excel_formula_map,
    excel_read_range,
    excel_workbook_overview,
    list_workpapers,
)

XLS = "구형조서_5300 현금.xls"
DOCX = "내부통제 기술서.docx"
XLSX = "일반조서.xlsx"


@pytest.fixture()
def doc_dir(tmp_path, monkeypatch):
    """xls(값·불리언·날짜·병합·숨김시트)·docx(제목·문단·표)·xlsx를 심은 폴더."""
    wb = xlwt.Workbook()
    ws = wb.add_sheet("잔액명세")
    ws.write(0, 0, "계정")
    ws.write(0, 1, "금액")
    ws.write(1, 0, "현금")
    ws.write(1, 1, 1500000.0)
    ws.write(2, 0, "확인여부")
    ws.write(2, 1, True)
    date_style = xlwt.easyxf(num_format_str="YYYY-MM-DD")
    ws.write(3, 0, "기준일")
    ws.write(3, 1, datetime.datetime(2026, 6, 30), date_style)
    ws.write_merge(5, 5, 0, 1, "병합 제목")
    hidden = wb.add_sheet("숨김시트")
    hidden.write(0, 0, "전기자료")
    hidden.visibility = 1
    wb.save(str(tmp_path / XLS))

    doc = DocxDocument()
    doc.add_heading("내부통제 기술서", level=1)
    doc.add_paragraph("판매 프로세스의 통제 활동을 기술한다.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "통제"
    table.cell(0, 1).text = "담당자"
    table.cell(1, 0).text = "출하 승인"
    table.cell(1, 1).text = "영업팀장"
    doc.save(str(tmp_path / DOCX))

    xwb = Workbook()
    xwb.active.cell(row=1, column=1, value="일반")
    xwb.save(str(tmp_path / XLSX))

    monkeypatch.setenv("WORKPAPERS_DIR", str(tmp_path))
    return tmp_path


def test_list_workpapers_includes_xls_and_docx(doc_dir):
    out = list_workpapers.invoke({})
    assert XLS in out
    assert DOCX in out
    assert XLSX in out


def test_xls_read_range_values(doc_dir):
    out = excel_read_range.invoke(
        {"path": XLS, "sheet": "잔액명세", "cell_range": "A1:B4"}
    )
    assert "현금" in out
    assert "1500000" in out  # 정수화된 값
    assert "True" in out
    assert "2026-06-30" in out
    assert ".xls 구형 형식" in out


def test_xls_overview_sheets_and_merge(doc_dir):
    out = excel_workbook_overview.invoke({"path": XLS})
    assert "잔액명세" in out
    assert "숨김시트" in out and "hidden" in out
    assert "병합 1건" in out


def test_xls_formula_map_notice(doc_dir):
    out = excel_formula_map.invoke({"path": XLS, "sheet": "잔액명세"})
    assert ".xls 구형 형식은 수식을 제공하지 않습니다" in out


def test_read_document_docx(doc_dir):
    out = read_document.invoke({"path": DOCX})
    assert out.startswith(f"문서: {DOCX}")
    assert "# 내부통제 기술서" in out
    assert "판매 프로세스의 통제 활동을 기술한다." in out
    assert "| 출하 승인 | 영업팀장 |" in out


def test_read_document_rejects_excel(doc_dir):
    out = read_document.invoke({"path": XLSX})
    assert out.startswith("오류:")
    assert "excel_workbook_overview" in out


def test_missing_file_listing_shows_docx(doc_dir):
    out = excel_read_range.invoke(
        {"path": "없는파일.xlsx", "sheet": "A", "cell_range": "A1"}
    )
    assert "파일이 없습니다" in out
    assert DOCX in out
